#!/usr/bin/env python3
"""Genera GUIOSPRO_Guia_Tecnica.docx desde la guía del proyecto."""

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Instale: pip install python-docx")
    raise SystemExit(1)

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "GUIOSPRO_Guia_Tecnica.docx"


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold:
        r.bold = True
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor(55, 65, 81)


def code_block(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("GUIOSPRO_FLOSS\n")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(15, 118, 110)
    t.add_run("\nGuía técnica del sistema\n")
    t.add_run("Por qué no se usa HTML · Arquitectura · Lógica por archivo\n\n")
    t.add_run("Método GUIOSAD v0.1.2\n")
    doc.add_page_break()

    heading(doc, "1. Por qué no usamos HTML como interfaz principal", 1)

    heading(doc, "1.1 Qué había al inicio del proyecto", 2)
    para(doc, "La primera versión del software utilizaba:")
    table(
        doc,
        ["Componente", "Función"],
        [
            ("Flexx (main.py)", "Framework de interfaz en Python que generaba UI en el navegador"),
            ("guiosad.html", "Archivo HTML muy grande (exportación / vista legada)"),
            ("CSV (factors.csv, guiosad_data.csv)", "Catálogo GUIOSAD sin base de datos"),
        ],
    )
    para(
        doc,
        "Era una aplicación con muchos formularios (factores, subfactores, escalas 1–4, "
        "FODA, recomendaciones). Mantener eso en HTML estático y lógica separada complicaba "
        "correcciones y evolución.",
    )

    heading(doc, "1.2 Qué se usa ahora", 2)
    table(
        doc,
        ["Antes", "Ahora"],
        [
            ("HTML estático + Flexx", "Streamlit (app.py)"),
            ("Archivos .html en el repositorio", "Python describe pantallas; Streamlit genera HTML al ejecutar"),
            ("Sin persistencia unificada", "SQLite + capas db/ y services/"),
        ],
    )

    heading(doc, "1.3 Razones técnicas de la decisión", 2)
    numbered(
        doc,
        [
            "Muchos formularios y pasos — Login, historial, pasos 1–6, borrador, informes. "
            "En Streamlit cada control (st.form, st.selectbox, st.slider) se define en pocas líneas de Python.",
            "Un solo lenguaje — Pantalla, reglas GUIOSAD y base de datos en Python.",
            "Menos errores de integración — Separar engine.py (lógica) de app.py (UI) clarifica el método GUIOSAD.",
            "Despliegue sencillo — streamlit run app.py y Streamlit Community Cloud.",
            "Misma experiencia en navegador — El usuario ve una aplicación web; Streamlit escribe el HTML.",
        ],
    )

    heading(doc, "1.4 ¿Significa que no hay HTML en absoluto?", 2)
    para(doc, "No. El navegador siempre renderiza HTML. En este proyecto:")
    bullets(
        doc,
        [
            "Streamlit genera la estructura de la página automáticamente.",
            "ui/styles.py aporta CSS (apariencia).",
            "En app.py y ui/components.py hay fragmentos HTML pequeños (login, tarjetas) "
            "mediante st.markdown(..., unsafe_allow_html=True). Son detalles visuales.",
        ],
    )

    heading(doc, "1.5 Conclusión para documentación académica", 2)
    quote(
        doc,
        "La interfaz de GUIOSPRO_FLOSS se implementó con el framework Streamlit, que construye "
        "dinámicamente la capa de presentación en el navegador a partir de código Python, en lugar "
        "de mantener plantillas HTML estáticas. Esta decisión responde a la complejidad del flujo de "
        "evaluación (múltiples formularios y pasos del método GUIOSAD), a la integración con la base "
        "de datos y a la simplificación del despliegue y mantenimiento del sistema.",
    )

    heading(doc, "2. Arquitectura general (capas)", 1)
    code_block(
        doc,
        """
NAVEGADOR (HTML/CSS/JS generado por Streamlit)
        |
PRESENTACIÓN — app.py, ui/components.py, ui/styles.py
        |
AUTENTICACIÓN — auth/service.py
        |
LÓGICA GUIOSAD — engine.py + guiosad.py
        |
SERVICIOS — evaluation_repo, catalog, export_report, audit
        |
PERSISTENCIA — db/models, session, config, bootstrap
        |
ALMACENAMIENTO — SQLite (data/guiospro.db)
        """,
    )

    heading(doc, "3. Flujo principal de la aplicación", 1)
    numbered(
        doc,
        [
            "El usuario ejecuta: streamlit run app.py",
            "main() llama a ensure_database_ready() — tablas, CSV, usuarios demo.",
            "Sin sesión → pantalla de login → authenticate().",
            "Tras login → Dashboard o Historial.",
            "Crear o abrir evaluación → EvaluationEngine desde BD (engine_from_evaluation).",
            "Pasos 1–2 — factores → guardar borrador → save_engine_to_db.",
            "Pasos 3–4 — subfactores → ponderación y FODA.",
            "Pasos 5–6 — FODA, completar, recomendación A/B/C → export_pdf / export_excel.",
        ],
    )

    heading(doc, "4. Descripción y lógica de cada archivo", 1)

    heading(doc, "4.1 Raíz del proyecto", 2)
    table(
        doc,
        ["Archivo", "Qué hace", "Lógica principal"],
        [
            (
                "app.py",
                "Punto de entrada Streamlit",
                "Orquesta pantallas, sesión, navegación; guardar y exportar; caché del motor.",
            ),
            (
                "engine.py",
                "Cerebro GUIOSAD",
                "Importancia relativa, FODA, recomendaciones A/B/C; sin UI ni SQL directo.",
            ),
            (
                "guiosad.py",
                "Catálogo desde CSV",
                "Dimensiones, 18 factores, 61 subfactores; etiquetas del método.",
            ),
            ("factors.csv", "Datos de factores", "Importancia sugerida y alcance."),
            ("guiosad_data.csv", "Datos de subfactores", "Textos de criterios por factor."),
            ("requirements.txt", "Dependencias", "streamlit, sqlalchemy, pandas, bcrypt, etc."),
            (".env", "Configuración", "GUIOSPRO_DB_MODE=local → SQLite."),
            (".streamlit/config.toml", "Config Streamlit", "Tema y servidor."),
        ],
    )

    heading(doc, "4.2 Carpeta auth/", 2)
    table(
        doc,
        ["Archivo", "Qué hace", "Lógica principal"],
        [
            (
                "auth/service.py",
                "Login y contraseñas",
                "authenticate() con bcrypt; AuthUser y permisos por rol.",
            ),
        ],
    )

    heading(doc, "4.3 Carpeta db/", 2)
    table(
        doc,
        ["Archivo", "Qué hace", "Lógica principal"],
        [
            ("db/config.py", "Conexión BD", "SQLite en data/guiospro.db (modo local)."),
            ("db/models.py", "ORM", "9 tablas: organizaciones, usuarios, catálogo, evaluaciones, auditoría."),
            ("db/session.py", "Sesiones SQLAlchemy", "Motor, get_session(), init_tables()."),
            (
                "db/bootstrap.py",
                "Arranque",
                "ensure_database_ready(): tablas, catálogo CSV, usuarios demo.",
            ),
        ],
    )

    heading(doc, "4.4 Carpeta services/", 2)
    table(
        doc,
        ["Archivo", "Qué hace", "Lógica principal"],
        [
            ("services/catalog.py", "Catálogo en BD", "CSV → dimensiones, factores, subfactores."),
            (
                "services/evaluation_repo.py",
                "Evaluaciones",
                "Crear, listar, cargar/guardar motor; DbGuiosadModel.",
            ),
            ("services/export_report.py", "Informes", "PDF (reportlab) y Excel (openpyxl)."),
            ("services/audit.py", "Auditoría", "log_action() en tabla auditoria."),
        ],
    )

    heading(doc, "4.5 Carpeta ui/", 2)
    table(
        doc,
        ["Archivo", "Qué hace", "Lógica principal"],
        [
            ("ui/styles.py", "CSS", "CUSTOM_CSS: colores, tarjetas, sidebar."),
            ("ui/components.py", "Componentes", "KPIs, sidebar, evaluaciones recientes."),
        ],
    )

    heading(doc, "4.6 Generado en ejecución", 2)
    table(
        doc,
        ["Ruta", "Qué es"],
        [
            ("data/guiospro.db", "Base SQLite (no subir a GitHub)."),
            ("__pycache__/", "Bytecode temporal."),
        ],
    )

    heading(doc, "5. Lógica del método GUIOSAD (engine.py)", 1)
    table(
        doc,
        ["Paso", "Lógica"],
        [
            ("Importancia del decisor", "Escala 1–4 por factor."),
            ("Importancia relativa", "Media sugerida + decisor → Irrelevante…Fundamental."),
            ("Factor relevante", "Si no es Irrelevante, se evalúan subfactores."),
            ("Subfactores", "Escala 1–4; ponderación = media."),
            ("FODA", "Interno ≥3 Fortaleza, <3 Debilidad; Externo Oportunidad/Amenaza."),
            ("Recomendación", "Reglas A, B, C según FODA."),
        ],
    )

    heading(doc, "6. Pantallas de app.py", 1)
    table(
        doc,
        ["Función", "Pantalla"],
        [
            ("page_login", "Acceso usuario/contraseña"),
            ("page_dashboard", "KPIs y resumen"),
            ("page_historial", "Lista, nueva evaluación, re-evaluación"),
            ("page_factores", "Pasos 1–2"),
            ("page_subfactores", "Pasos 3–4"),
            ("page_foda", "Pasos 5–6 e informes"),
            ("render_sidebar", "Menú lateral"),
            ("main", "Enrutador principal"),
        ],
    )

    heading(doc, "7. Resumen: archivo → responsabilidad", 1)
    table(
        doc,
        ["Archivo", "Responsabilidad en una línea"],
        [
            ("app.py", "Pantallas y conexión con servicios"),
            ("engine.py", "Cálculo GUIOSAD, FODA y recomendación"),
            ("guiosad.py", "Catálogo desde CSV"),
            ("auth/service.py", "Login y roles"),
            ("db/config.py", "Ruta a SQLite"),
            ("db/models.py", "Definición de tablas"),
            ("db/session.py", "Conexión BD"),
            ("db/bootstrap.py", "Inicialización al arrancar"),
            ("services/catalog.py", "CSV → BD"),
            ("services/evaluation_repo.py", "Guardar/cargar evaluaciones"),
            ("services/export_report.py", "PDF y Excel"),
            ("services/audit.py", "Registro de acciones"),
            ("ui/styles.py", "Diseño visual (CSS)"),
            ("ui/components.py", "Bloques de UI"),
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GUIOSPRO_FLOSS — Método GUIOSAD v0.1.2 — Documento técnico")
    r.italic = True
    r.font.size = Pt(10)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Documento generado: {path}")
